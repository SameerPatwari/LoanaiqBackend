import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# Load CSV into DataFrame
csv_file = 'statement_of_profitability.csv'
df = pd.read_csv(csv_file)

# Create a new Document
doc = Document()

# Add header
header = doc.sections[0].header
header_table = header.add_table(rows=1, cols=2, width=Inches(6))
header_table.autofit = False

# Set column widths
header_table.columns[0].width = Inches(2)
header_table.columns[1].width = Inches(4)

# Add logo
logo_cell = header_table.cell(0, 0)
logo_run = logo_cell.paragraphs[0].add_run()
logo_run.add_picture('static/images/logo.png', width=Inches(1.5))
logo_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add business name and details to the right
text_cell = header_table.cell(0, 1)
text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
text = text_cell.paragraphs[0].add_run("JANAKALYAN SAHAKARI BANK LTD.\nCredit Department\nNote To BLSC")
text.bold = True
text_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add table to the document
doc.add_paragraph("\nFinancial Position: (in lakhs)", style='Heading1')

table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'

# Add column headers
hdr_cells = table.rows[0].cells
for idx, column in enumerate(df.columns):
    hdr_cells[idx].text = column

# Add rows to the table
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    for idx, value in enumerate(row):
        row_cells[idx].text = f'{value}'

# Add text from the text file
with open('uploads/gpt_response.txt', 'r') as file:
    text_content = file.read()

doc.add_paragraph("\n" + text_content)

# Add footer with separation line and page number
footer = doc.sections[0].footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.add_run("___________________________________").add_break()
footer_paragraph = footer.add_paragraph()
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

pager = footer_paragraph.add_run()
field_code = OxmlElement('w:fldSimple')
field_code.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr', 'PAGE')
pager._element.append(field_code)

# Save the document
doc.save('output.docx')