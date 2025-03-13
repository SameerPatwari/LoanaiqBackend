from fastapi import FastAPI, UploadFile, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from openai import OpenAI
import boto3
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import PyPDF2
import glob

# Load environment variables
load_dotenv()

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize AWS S3 client
s3 = boto3.client(
    's3',
    aws_access_key_id=os.getenv('AWS_ACCESS_KEY'),
    aws_secret_access_key=os.getenv('AWS_SECRET_KEY')
)

# Initialize OpenAI client
openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Create necessary directories if they don't exist
UPLOADS_DIR = "uploads"
OUTPUT_DIR = "output"
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

class UserIdRequest(BaseModel):
    user_id: str

class NoteRequest(BaseModel):
    user_id: str

@app.post("/load-data")
async def load_data(request: UserIdRequest):
    try:
        # Construct filename
        filename = f"{request.user_id}.json"
        
        # Download file from S3
        local_file_path = os.path.join(UPLOADS_DIR, filename)
        s3.download_file(
            os.getenv('S3_BUCKET_NAME'),
            filename,
            local_file_path
        )
        
        return JSONResponse(content={"status": "success", "message": "Data loaded successfully"})
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/upload-profile")
async def upload_profile(file: UploadFile):
    try:
        # Save file locally
        file_path = os.path.join(UPLOADS_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return JSONResponse(content={"status": "success", "message": "Profile uploaded successfully"})
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

def extract_text_from_document(file_path):
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            return ' '.join([page.extract_text() for page in pdf_reader.pages])
    return ""

def filter_data_for_prompt(data, prompt_type):
    # Load ratios configuration
    with open('ratios.json', 'r') as f:
        ratio_config = json.load(f)
    
    # Filter data based on prompt type and ratio configuration
    filtered_data = {}
    if prompt_type in ratio_config:
        config = ratio_config[prompt_type]
        for category in config:
            if category in data:
                filtered_data[category] = {
                    k: v for k, v in data[category].items()
                    if k in config[category]
                }
    
    return filtered_data

def filter_data_for_table(data, prompt_type):
    # Load ratios configuration
    with open('tables.json', 'r') as f:
        ratio_config = json.load(f)
    
    # Filter data based on prompt type and ratio configuration
    filtered_data = {}
    if prompt_type in ratio_config:
        config = ratio_config[prompt_type]
        for category in config:
            if category in data:
                filtered_data[category] = {
                    k: v for k, v in data[category].items()
                    if k in config[category]
                }
    
    return filtered_data

def generate_document(user_data, response_text, user_id):
    doc = Document()
    
    # Set narrow margins (0.5 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.left_margin = int(Inches(0.5))
        section.right_margin = int(Inches(0.5))
        section.top_margin = int(Inches(0.5))
        section.bottom_margin = int(Inches(0.5))

    # Add cover page with APPRAISAL NOTE
    cover_heading = doc.add_heading('APPRAISAL NOTE', level=0)
    cover_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_heading.style.font.size = Pt(20)
    cover_heading.style.font.bold = True
    
    # Add index
    index_heading = doc.add_paragraph()
    index_run = index_heading.add_run("\n\nINDEX")
    index_run.bold = True
    index_run.font.size = Pt(14)
    
    # Add index entries placeholder
    index_entries = []  # Store entries to add later
    
    # Add page break after index
    doc.add_page_break()
    
    # Add header to all pages except cover
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=Inches(7.5))
    header_table.autofit = False

    # Set column widths
    header_table.columns[0].width = int(Inches(2.5))
    header_table.columns[1].width = int(Inches(5.0))

    # Add logo
    logo_cell = header_table.cell(0, 0)
    logo_run = logo_cell.paragraphs[0].add_run()
    logo_run.add_picture('static/images/logo.png', width=Inches(2.0))
    logo_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add business name and details to the right
    text_cell = header_table.cell(0, 1)
    text_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text = text_cell.paragraphs[0].add_run("JANAKALYAN SAHAKARI BANK LTD.\nCredit Department\nNote To BLSC")
    text.bold = True
    text.font.size = Pt(12)
    text_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space after header
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # List to store all analyses for recommendations
    all_field_analyses = []
    
    # Dictionary to store section bookmarks
    bookmarks = {}
    
    def add_bookmark(paragraph, bookmark_id):
        """Add a bookmark to a paragraph"""
        # Create bookmark start and end elements
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_id)
        paragraph._p.append(start)
        
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), bookmark_id)
        paragraph._p.append(end)
    
    # Add company metadata
    about_company = user_data.get('metadata', {}).get('about_company', '')
    if about_company:
        heading = doc.add_heading('About the Company', level=1)
        heading.style.font.size = Pt(16)
        heading.style.font.bold = True
        add_bookmark(heading, 'about_company')
        para = doc.add_paragraph(about_company)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add main analysis heading
    analysis_heading = doc.add_heading('Analysis Of Financial Statement', level=1)
    analysis_heading.style.font.size = Pt(16)
    analysis_heading.style.font.bold = True
    add_bookmark(analysis_heading, 'analysis')

    # Process and format the response text
    lines = response_text.strip().split('\n')
    current_section = None
    current_analysis = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line in ["Ratios Analysis:", "Balance Sheet Analysis:", "Profit and Loss Analysis:"]:
            # Main heading
            section_name = line.strip(':').lower().replace(' ', '_')
            heading = doc.add_heading('', level=2)
            run = heading.add_run(line.strip(':'))
            run.font.size = Pt(14)
            run.font.bold = True
            add_bookmark(heading, section_name)
            current_section = line.split()[0].lower()
        elif line.endswith("Table"):
            # Add appropriate table based on current section
            section_key = current_section
            if section_key == "balance":
                section_key = "balance_sheet"
            elif section_key == "profit":
                section_key = "profit_loss"
            
            if section_key in user_data:
                print(f"Generating table for {section_key}")
                # Remove the word "Table" from the heading when adding the table
                table_name = line.replace(" Table", "")
                add_table_to_document(doc, table_name, user_data[section_key])
        elif line.startswith("Analysis of"):
            # Store previous analysis if exists
            if current_analysis:
                all_field_analyses.append("\n".join(current_analysis))
                current_analysis = []
            # Start new analysis
            current_analysis = [line]
            # Subheading
            heading = doc.add_heading('', level=3)
            run = heading.add_run(line)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue color
        else:
            # Add line to current analysis
            current_analysis.append(line)
            # Content bullet points
            if line.startswith('- '):
                line = line[2:]
            if line.startswith('**') and line.endswith('**'):
                line = line[2:-2]
            
            paragraph = doc.add_paragraph()
            paragraph.style.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if line.startswith("Financial Risk:"):
                # Add bullet point
                bullet_run = paragraph.add_run('• ')
                bullet_run.font.size = Pt(11)
                
                # Split into label and content
                risk_parts = line.split(":", 1)
                risk_label = risk_parts[0] + ":"
                risk_content = risk_parts[1].strip() if len(risk_parts) > 1 else ""
                
                # Add "Financial Risk:" in red
                risk_run = paragraph.add_run(risk_label + " ")
                risk_run.font.size = Pt(11)
                risk_run.font.bold = True
                risk_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Dark red color
                
                # Add the risk content
                if risk_content:
                    content_run = paragraph.add_run(risk_content)
                    content_run.font.size = Pt(11)
            else:
                # Add bullet point
                bullet_run = paragraph.add_run('• ')
                bullet_run.font.size = Pt(11)
                # Add the content
                content_run = paragraph.add_run(line)
                content_run.font.size = Pt(11)
            
            # Set consistent paragraph formatting
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Store last analysis if exists
    if current_analysis:
        all_field_analyses.append("\n".join(current_analysis))

    # Add Recommendations section
    recommendations_heading = doc.add_heading('Recommendations', level=2)
    recommendations_heading.style.font.size = Pt(14)
    add_bookmark(recommendations_heading, 'recommendations')

    # Generate recommendations using GPT
    try:
        # Combine all analyses and truncate to fit token limit
        all_analyses_text = "\n\n".join(all_field_analyses)
        # Roughly estimate tokens (4 chars per token) and limit to 6000 tokens
        max_chars = 24000  # 6000 tokens * 4 chars per token
        if len(all_analyses_text) > max_chars:
            all_analyses_text = all_analyses_text[:max_chars] + "\n[Analysis truncated due to length...]"
        
        # Load recommendations prompt
        with open(os.path.join("prompt_library", "recommendations.txt"), 'r', encoding='utf-8') as f:
            recommendations_prompt = f.read()
        
        # Get recommendations from GPT
        response = openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": recommendations_prompt},
                {"role": "user", "content": f"Based on the following analyses, provide recommendations:\n\n{all_analyses_text}"}
            ],
            temperature=0,
            top_p=0.1
        )
        
        # Add recommendations to document
        recommendations = response.choices[0].message.content.strip()
        for line in recommendations.split('\n'):
            if line.strip():
                if line.startswith('#') or line.startswith('*'):
                    # Handle headers or bullet points
                    para = doc.add_paragraph(line.lstrip('#* '))
                    para.style.font.size = Pt(11)
                else:
                    para = doc.add_paragraph(line)
                    para.style.font.size = Pt(11)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception as e:
        print(f"Error generating recommendations: {str(e)}")
        # Add error message to document
        doc.add_paragraph("Unable to generate recommendations at this time.")

    # Update index with hyperlinks
    index_items = [
        ("About the Company", 1, 'about_company'),
        ("Analysis Of Financial Statement", 1, 'analysis'),
        ("Ratios Analysis", 2, 'ratios_analysis'),
        ("Balance Sheet Analysis", 2, 'balance_sheet_analysis'),
        ("Profit and Loss Analysis", 2, 'profit_and_loss_analysis'),
        ("Recommendations", 2, 'recommendations')
    ]
    
    for title, level, bookmark_id in index_items:
        # Create paragraph for index entry
        index_para = doc.add_paragraph()
        index_para.paragraph_format.left_indent = Inches(0.5 if level == 2 else 0)
        
        # Add hyperlink
        hyperlink = add_hyperlink(doc, index_para, title, bookmark_id)
        hyperlink.font.size = Pt(11)
        if level == 1:
            hyperlink.font.bold = True
        
        # Add dot leaders
        dots = index_para.add_run('.' * 50)
        dots.font.size = Pt(11)
        
        # Add newline
        index_para.paragraph_format.space_after = Pt(6)

    # Save document
    output_path = os.path.join(OUTPUT_DIR, f'analysis_{user_id}.docx')
    doc.save(output_path)
    return output_path

def add_hyperlink(doc, paragraph, text, bookmark_id):
    """
    Add a hyperlink to a paragraph that links to a bookmark
    """
    # Create the w:hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_id)
    
    # Create a new run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink style
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0000FF")
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), "single")
    rPr.append(u)
    
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    
    r = paragraph.add_run()
    r._r.append(hyperlink)
    return r

def qn(tag):
    """
    Turn a namespace prefixed tag into a Clark-notation qualified tag
    """
    prefix, tagname = tag.split(':')
    uri = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }[prefix]
    return '{{{}}}{}'.format(uri, tagname)

def add_table_to_document(doc, table_name, table_data):
    """Add a table to the document with the given name and data"""
    # Add table heading
    heading = doc.add_heading('', level=2)
    run = heading.add_run(table_name)
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Get years and data
    years = table_data['years']
    
    # Create table with years as columns
    table = doc.add_table(rows=1, cols=len(years) + 1)
    table.style = 'Table Grid'
    table.allow_autofit = False
    
    # Calculate available width
    available_width = int(Inches(7.5))
    first_col_width = int(Inches(3.0))
    remaining_width = available_width - first_col_width
    year_col_width = int(remaining_width / len(years))
    
    # Apply column widths
    table.columns[0].width = first_col_width
    for i in range(1, len(table.columns)):
        table.columns[i].width = year_col_width
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Financial Term'
    for idx, year in enumerate(years):
        header_cells[idx + 1].text = year
    
    # Format header row
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
        run.font.bold = True
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for category, items in table_data.items():
        if category != 'years':
            for item_name, values in items.items():
                row_cells = table.add_row().cells
                row_cells[0].text = item_name
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for idx, value in enumerate(values):
                    cell = row_cells[idx + 1]
                    cell.text = str(value)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Format all cells
    for row in table.rows:
        row.height = int(Inches(0.25))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # Add space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def cleanup_user_files(userid):
    """Delete all files related to the userid from uploads and output folders"""
    # Clean uploads folder
    upload_pattern = os.path.join(UPLOADS_DIR, f"{userid}*")
    for file in glob.glob(upload_pattern):
        os.remove(file)
    
    # Clean output folder
    output_pattern = os.path.join(OUTPUT_DIR, f"{userid}*")
    for file in glob.glob(output_pattern):
        os.remove(file)

@app.post("/generate-note")
async def generate_note(request: NoteRequest):
    try:
        # Check if user data exists, if not try to load it
        user_data_path = os.path.join(UPLOADS_DIR, f"{request.user_id}.json")
        if not os.path.exists(user_data_path):
            try:
                # Try to download from S3
                s3.download_file(
                    os.getenv('S3_BUCKET_NAME'),
                    f"{request.user_id}.json",
                    user_data_path
                )
            except Exception as s3_error:
                raise HTTPException(
                    status_code=404,
                    detail=f"User data not found. Please ensure data is loaded first. Error: {str(s3_error)}"
                )

        # Load user data for document generation
        try:
            with open(user_data_path, 'r', encoding='utf-8') as f:
                user_data = json.load(f)
        except json.JSONDecodeError as json_error:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid JSON data format. Error: {str(json_error)}"
            )
        except Exception as read_error:
            raise HTTPException(
                status_code=500,
                detail=f"Error reading user data: {str(read_error)}"
            )

        # Process each section (ratios, balance_sheet, profit_loss)
        sections = [
            {
                'name': 'Ratios',
                'data_key': 'ratios',
                'prompt_prefix': 'ratios',
                'heading': 'Ratios Analysis:'
            },
            {
                'name': 'Balance Sheet',
                'data_key': 'balance_sheet',
                'prompt_prefix': 'balance_sheet',
                'heading': 'Balance Sheet Analysis:'
            },
            {
                'name': 'Profit and Loss',
                'data_key': 'profit_loss',
                'prompt_prefix': 'profit_loss',
                'heading': 'Statement Of Profitability Analysis:'
            }
        ]

        # Initialize list to store all analyses
        all_analyses = []

        for section in sections:
            print(f"\nProcessing {section['name']} section")
            
            # Get section data
            section_data = user_data.get(section['data_key'])
            if not section_data:
                print(f"No {section['name']} data found in user data")
                continue

            # Add section heading
            all_analyses.append(section['heading'])
            
            print(f"Available categories in {section['name']}:", list(section_data.keys()))

            # Add table for this section (using section_data directly instead of filtered data)
            if section_data and 'years' in section_data:  # If section has data and years, it can be displayed as a table
                print(f"Adding table marker for {section['name']}")
                all_analyses.append(f"{section['name']} Table")  # This will trigger table generation in the document

            # Process each field in the section
            for category, fields in section_data.items():
                if category == 'years':  # Skip the years array
                    continue
                
                print(f"\nProcessing category in {section['name']}: {category}")
                print(f"Fields in {category}:", list(fields.keys()))
                
                for field in fields.keys():
                    # Construct prompt filename
                    prompt_file = f"{section['prompt_prefix']}_{category}_{field}.txt"
                    prompt_path = os.path.join("prompt_library", prompt_file)
                    
                    print(f"Looking for prompt file: {prompt_path}")
                    
                    if os.path.exists(prompt_path):
                        print(f"Found prompt file for {field}")
                        # Load prompt
                        with open(prompt_path, 'r', encoding='utf-8') as f:
                            prompt_content = f.read()
                        
                        # Get field data
                        years = section_data['years']
                        values = section_data[category][field]
                        
                        # Create data string
                        data_str = "Years             " + "   ".join(years) + "\n\n"
                        data_str += f"{field}\t   " + "\t   ".join(str(val) for val in values)
                        
                        print(f"Data for {field}:")
                        print(data_str)
                        
                        # Replace placeholder in prompt with actual data
                        final_prompt = prompt_content.replace("DATA:", f"DATA:\n{data_str}")
                        
                        try:
                            # Get GPT analysis
                            response = openai_client.chat.completions.create(
                                model="gpt-4",
                                messages=[
                                    {"role": "system", "content": final_prompt},
                                    {"role": "user", "content": f"""Please analyze the following financial data and provide insights:

Years: {', '.join(years)}
Values: {', '.join(str(val) for val in values)}

Please format your response exactly as follows:

Analysis of {field}:
- [First insight following the format: The change 'X' in parameter 'Y' maybe caused by 'Z' and affects 'A']
- [Second insight following the format: The change 'X' in parameter 'Y' maybe caused by 'Z' and affects 'A']
- [Third insight following the format: The change 'X' in parameter 'Y' maybe caused by 'Z' and affects 'A']
- Financial Risk: [One comprehensive risk statement]"""}
                                ],
                                temperature=0,
                                top_p=0.1
                            )
                            
                            # Add analysis to list
                            analysis_text = response.choices[0].message.content.strip()
                            if not analysis_text.startswith("Analysis of"):
                                analysis_text = f"Analysis of {field}:\n" + analysis_text
                            all_analyses.append(analysis_text)
                            print(f"Successfully generated analysis for {field}")
                        except Exception as gpt_error:
                            print(f"Error getting GPT analysis for {field}: {str(gpt_error)}")
                            continue  # Skip this field and continue with others
                    else:
                        print(f"No prompt file found for {field}")

        print(f"\nTotal analyses generated: {len(all_analyses)}")
        
        if len(all_analyses) == 0:
            raise HTTPException(
                status_code=500,
                detail="No analyses were generated. Please check the prompt library and GPT service."
            )

        # Combine all analyses with proper spacing
        combined_analysis = "\n\n".join(all_analyses)
        
        # Save combined analysis
        response_path = os.path.join(OUTPUT_DIR, f"{request.user_id}_response.json")
        with open(response_path, 'w', encoding='utf-8') as f:
            json.dump({"financial_position": combined_analysis}, f)
        
        # Generate document
        doc_path = generate_document(user_data, combined_analysis, request.user_id)
        
        # After sending response, clean up files
        cleanup_user_files(request.user_id)
        
        return FileResponse(
            doc_path, 
            filename=f"analysis_{request.user_id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        print(f"Error in generate_note: {str(e)}")  # Add this for debugging
        # Clean up any partial files that might have been created
        cleanup_user_files(request.user_id)
        raise HTTPException(status_code=400, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)